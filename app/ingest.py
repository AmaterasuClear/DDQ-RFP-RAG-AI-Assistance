from __future__ import annotations

import hashlib
import re
import shutil
from pathlib import Path

from app.config import (
    CHUNK_SIZE,
    CHUNK_OVERLAP,
    SUPPORTED_DOCUMENT_EXTENSIONS,
    UPLOADS_DIR,
    ensure_runtime_dirs,
)
from app.embeddings import generate_embedding as _generate_embedding
from app.schemas import Chunk, DocumentRecord, utc_now_iso
from app.storage import audit_event, delete_document, save_chunks, save_document_record
from app.text_cleaner import clean_text, is_corrupted_chunk


# ── New chunk size: 800 chars with 100 overlap, word-boundary aware ────────────
_CHUNK_SIZE = 400
_CHUNK_OVERLAP = 60


# ── Document loaders ───────────────────────────────────────────────────────────


def _pdf_reader(path: Path):
    try:
        from PyPDF2 import PdfReader
    except ImportError as exc:
        raise RuntimeError("PyPDF2 is required for PDF ingestion") from exc
    return PdfReader(str(path))


def load_pdf(pdf_path: str | Path) -> list[tuple[int, str]]:
    path = Path(pdf_path)
    if not path.exists():
        raise FileNotFoundError(path)
    reader = _pdf_reader(path)
    pages: list[tuple[int, str]] = []
    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = clean_text(text)
        if text:
            pages.append((index, text))
    return pages


def load_docx(docx_path: str | Path) -> list[tuple[int, str]]:
    """Parse DOCX preserving paragraph structure.

    Each paragraph is extracted as a unit. Tables are extracted with
    cell content joined logically. The output preserves paragraph
    breaks as \n\n for downstream chunking to split on.
    """
    path = Path(docx_path)
    if not path.exists():
        raise FileNotFoundError(path)
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is required for DOCX ingestion") from exc

    document = Document(str(path))

    # Collect paragraphs preserving structure
    paragraphs: list[str] = []
    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        # Use paragraph style to detect headings
        style_name = (para.style.name if para.style else "").lower()
        if "heading" in style_name or "title" in style_name:
            text = f"\n{text}\n"
        paragraphs.append(text)

    # Collect tables with context
    for table in document.tables:
        table_lines: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                table_lines.append(" | ".join(cells))
        if table_lines:
            paragraphs.append("\n".join(table_lines))

    # Build document text with paragraph breaks preserved
    raw_text = "\n\n".join(paragraphs)

    # Clean the extracted text
    cleaned = clean_text(raw_text)
    if not cleaned:
        return []

    # Split into logical pages/sections by heading markers
    # We use paragraph double-newlines as section boundaries for better chunking
    return [(1, cleaned)]


def load_text(text_path: str | Path) -> list[tuple[int, str]]:
    path = Path(text_path)
    if not path.exists():
        raise FileNotFoundError(path)
    raw = path.read_text(encoding="utf-8")
    cleaned = clean_text(raw)
    return [(1, cleaned)] if cleaned else []


def load_document(document_path: str | Path) -> list[tuple[int, str]]:
    path = Path(document_path)
    extension = path.suffix.lower()
    if extension == ".pdf":
        return load_pdf(path)
    if extension == ".docx":
        return load_docx(path)
    if extension in {".txt", ".md"}:
        return load_text(path)
    raise ValueError(f"Unsupported document type: {extension}")


# ── Smart chunking — NEVER splits words ────────────────────────────────────────


def _strip_leading_fragment(chunk: str) -> str:
    """Remove a sentence fragment from the beginning of a chunk.

    If a chunk starts mid-sentence (lowercase word, or clearly a continuation),
    advance to the first full sentence start.
    """
    if not chunk:
        return chunk

    # If chunk starts with uppercase or number, it's already clean
    stripped = chunk.lstrip()
    if stripped and stripped[0].isupper():
        return chunk
    if stripped and stripped[0].isdigit():
        return chunk

    # Find the first sentence start: period/exclamation/question + space + uppercase
    match = re.search(r"[.!?]\s+[A-Z]", stripped)
    if match:
        # Start from the capital letter after the sentence end
        new_start = match.start() + len(match.group()) - 1
        cleaned = stripped[new_start:].strip()
        if len(cleaned) >= 30:
            return cleaned

    # Find first uppercase word boundary
    match = re.search(r"\s([A-Z][a-z])", stripped)
    if match:
        cleaned = stripped[match.start() + 1 :].strip()
        if len(cleaned) >= 30:
            return cleaned

    # Can't fix — return original but marked
    return chunk


def _find_split_point(text: str, target: int) -> int:
    """Find the best split point near `target` that falls on a natural boundary.

    Priority: paragraph break > sentence end > word boundary > whitespace.
    """
    window = min(150, target // 2)
    search_start = max(0, target - window)
    search_end = min(len(text), target + window)

    search_region = text[search_start:search_end]
    offset = search_start

    # 1. Prefer paragraph break
    for boundary in ["\n\n", "\n"]:
        pos = search_region.rfind(boundary, 0, target - offset + window // 2)
        if pos > window // 4:
            return offset + pos + len(boundary)

    # 2. Prefer sentence end (. ! ? followed by space and capital letter)
    for match in re.finditer(r"[.!?]\s+(?=[A-Z])", search_region):
        if abs((offset + match.end()) - target) < window:
            return offset + match.end()

    # 3. Prefer sentence end followed by any space
    for match in re.finditer(r"[.!?]\s+", search_region):
        if abs((offset + match.end()) - target) < window:
            return offset + match.end()

    # 4. Prefer comma or semicolon
    for match in re.finditer(r"[,;]\s+", search_region):
        if abs((offset + match.end()) - target) < window // 2:
            return offset + match.end()

    # 5. Fall back to the last space within the window
    last_space = search_region.rfind(" ", 0, target - offset)
    if last_space > window // 2:
        return offset + last_space

    # 6. Last resort: target itself
    return target


def split_chunks(
    text: str,
    chunk_size: int = _CHUNK_SIZE,
    overlap: int = _CHUNK_OVERLAP,
) -> list[str]:
    """Split text into chunks without ever splitting a word in the middle.

    Uses natural boundaries: paragraphs, sentences, then word boundaries.
    """
    text = text.strip()
    if not text:
        return []
    if len(text) <= chunk_size:
        return [text]

    chunks: list[str] = []
    start = 0

    while start < len(text):
        end = min(start + chunk_size, len(text))

        if end < len(text):
            # Find a natural split point
            split_at = _find_split_point(text, end)
            # Don't let splits create tiny fragments
            if split_at - start < 150:
                # Extend to the next natural boundary
                split_at = _find_split_point(text, end + 200)
            end = split_at

        chunk = text[start:end].strip()
        chunk = _strip_leading_fragment(chunk)

        if chunk and len(chunk) >= 30 and not is_corrupted_chunk(chunk):
            chunks.append(chunk)

        # Advance with overlap
        next_start = end - overlap
        if next_start <= start:
            next_start = end
        # Align next start to a word boundary
        if next_start < len(text):
            while next_start < len(text) and text[next_start] not in (" ", "\n"):
                next_start += 1
        start = next_start

    return chunks


# ── Embedding helper ───────────────────────────────────────────────────────────


def generate_embedding(text: str) -> list[float]:
    return _generate_embedding(text)


# ── File persistence ───────────────────────────────────────────────────────────


def save_chunk(chunk: Chunk, db_path: str | Path | None = None) -> None:
    from app.storage import save_chunk as persist_chunk
    persist_chunk(chunk, db_path)


def persist_uploaded_pdf(source_path: str | Path) -> Path:
    ensure_runtime_dirs()
    source = Path(source_path)
    destination = UPLOADS_DIR / source.name
    if source.resolve() != destination.resolve():
        shutil.copy2(source, destination)
    return destination


def file_sha256(path: str | Path) -> str:
    digest = hashlib.sha256()
    with Path(path).open("rb") as file:
        for block in iter(lambda: file.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def persist_uploaded_document(source_path: str | Path) -> Path:
    ensure_runtime_dirs()
    source = Path(source_path)
    destination = UPLOADS_DIR / source.name
    if source.resolve() != destination.resolve():
        shutil.copy2(source, destination)
    return destination


# ── Main ingestion ─────────────────────────────────────────────────────────────


def ingest_document(
    document_path: str | Path,
    doc_type: str = "general",
    version: str = "v1",
    db_path: str | Path | None = None,
    copy_to_uploads: bool = True,
    replace_existing: bool = True,
) -> list[Chunk]:
    source_path = Path(document_path)
    if source_path.suffix.lower() not in SUPPORTED_DOCUMENT_EXTENSIONS:
        raise ValueError(f"Unsupported document type: {source_path.suffix}")

    stored_path = persist_uploaded_document(source_path) if copy_to_uploads else source_path
    if replace_existing:
        delete_document(stored_path.name, db_path)

    pages = load_document(stored_path)
    upload_date = utc_now_iso()
    chunks: list[Chunk] = []

    for page_number, page_text in pages:
        for text in split_chunks(page_text):
            if is_corrupted_chunk(text):
                continue
            chunks.append(
                Chunk(
                    doc_name=stored_path.name,
                    page=page_number,
                    text=text,
                    embedding=generate_embedding(text),
                    doc_type=doc_type,
                    version=version,
                    upload_date=upload_date,
                )
            )

    save_chunks(chunks, db_path)
    stat = stored_path.stat()
    save_document_record(
        DocumentRecord(
            doc_name=stored_path.name,
            source_path=str(stored_path.resolve()),
            file_type=stored_path.suffix.lower().lstrip("."),
            doc_type=doc_type,
            version=version,
            chunk_count=len(chunks),
            content_hash=file_sha256(stored_path),
            source_modified_at=utc_now_iso()
            if not stat.st_mtime
            else __import__("datetime").datetime.fromtimestamp(
                stat.st_mtime,
                tz=__import__("datetime").timezone.utc,
            ).isoformat(),
            indexed_at=upload_date,
        ),
        db_path,
    )
    audit_event(
        "document_ingested",
        {
            "doc_name": stored_path.name,
            "doc_type": doc_type,
            "version": version,
            "pages": len(pages),
            "chunks": len(chunks),
        },
        db_path,
    )
    return chunks


def ingest_pdf(
    pdf_path: str | Path,
    doc_type: str = "general",
    version: str = "v1",
    db_path: str | Path | None = None,
    copy_to_uploads: bool = True,
) -> list[Chunk]:
    return ingest_document(
        pdf_path,
        doc_type=doc_type,
        version=version,
        db_path=db_path,
        copy_to_uploads=copy_to_uploads,
    )